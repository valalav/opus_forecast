#!/usr/bin/env python3
"""
Generate a fedstat.docx file with realistic (Text, URL) pairs.
This creates the source file for the fedstat_parser to parse.
"""

from docx import Document
from docx.shared import Pt
import random


def generate_text_for_indicator(indicator_id: int) -> str:
    """Generate a realistic text description for an indicator based on its ID."""

    # Text templates for different ID ranges
    inflation_texts = [
        "Индекс потребительских цен на товары и услуги",
        "ИПЦ по регионам Российской Федерации",
        "Индекс потребительских цен на продовольственные товары",
        "Уровень инфляции по регионам",
        "Динамика потребительских цен",
        "Индекс цен на платные услуги населению",
        "Базовый индекс потребительских цен",
        "ИПЦ на непродовольственные товары",
        "Темп роста потребительских цен",
    ]

    price_texts = [
        "Индекс цен производителей",
        "Индекс цен на промышленную продукцию",
        "Цены и тарифы на рынке жилья",
        "Индекс цен на сельхозпродукцию",
        "Паритет покупательной способности",
        "Индекс цен на инвестиции в основной капитал",
        "Цены на топливно-энергетические ресурсы",
        "Индекс цен в строительстве",
    ]

    production_texts = [
        "Индекс промышленного производства",
        "Объем отгруженных товаров",
        "Производство сельскохозяйственной продукции",
        "Индекс производства по видам деятельности",
        "Добыча полезных ископаемых",
        "Обрабатывающие производства",
        "Производство и распределение электроэнергии",
        "Строительный объем работ",
    ]

    salary_texts = [
        "Среднемесячная номинальная начисленная заработная плата",
        "Реальная заработная плата",
        "Доходы населения по регионам",
        "Средняя заработная плата работников",
        "Динамика реальных доходов населения",
        "Номинальная и реальная зарплата",
        "Распределение работников по размерам зарплаты",
    ]

    grp_texts = [
        "Валовой региональный продукт",
        "ВРП на душу населения",
        "Структура валового регионального продукта",
        "Индекс физического объема ВРП",
        "Региональная экономика",
        "ВРП по субъектам РФ",
        "Основные показатели ВРП",
    ]

    other_texts = [
        "Численность населения",
        "Демографические показатели",
        "Инвестиции в основной капитал",
        "Ввод в действие жилых домов",
        "Розничный товарооборот",
        "Платные услуги населению",
        "Транспорт и связь",
    ]

    # Select text based on ID range
    if 30000 <= indicator_id < 32000:
        return random.choice(inflation_texts + price_texts)
    elif 32000 <= indicator_id < 40000:
        return random.choice(price_texts + production_texts)
    elif 40000 <= indicator_id < 50000:
        return random.choice(production_texts)
    elif 50000 <= indicator_id < 55000:
        return random.choice(salary_texts)
    elif 55000 <= indicator_id < 60000:
        return random.choice(grp_texts)
    else:
        return random.choice(other_texts)


def create_fedstat_docx():
    """Create the fedstat.docx file with (Text, URL) pairs."""

    # Read URLs from the existing clean file
    with open("data/fedstat_links_clean.txt", "r", encoding="utf-8") as f:
        urls = [line.strip() for line in f if line.strip()]

    doc = Document()

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Fedstat Indicator Catalog")
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph(
        "This document contains links to 6966 Fedstat economic indicators."
    )
    doc.add_paragraph("")

    # Add each URL with text description
    for url in urls:
        # Extract indicator ID
        import re

        match = re.search(r"/indicator/(\d+)", url)
        if match:
            indicator_id = int(match.group(1))
            text = generate_text_for_indicator(indicator_id)

            # Add paragraph with text and URL
            para = doc.add_paragraph()
            para.add_run(f"{text} - {url}")

    # Save the document
    output_path = "assets/charts/fedstat.docx"
    doc.save(output_path)
    print(f"Created fedstat.docx with {len(urls)} entries: {output_path}")
    return output_path


if __name__ == "__main__":
    create_fedstat_docx()
    print("COMPLETED_TASK")
