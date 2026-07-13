#!/usr/bin/env python3
"""
Fedstat Smart Link Prioritization

This agent implements a safe strategy to filter fedstat data series:
1. Parse the fedstat URL list from docx using python-docx
2. Create catalog CSV with >5000 name-link pairs
3. Filter by keywords (Price, Inflation, Production, GRP, Salary)
4. Fetch metadata for top candidates to verify data availability
5. Output: prioritized_fedstat_feed.csv

SAFETY: Only fetches metadata for top-n candidates (default 50) to avoid API bans.
Uses rate limiting and retry logic for robust HTTP requests.
"""

import re
import csv
import time
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from docx import Document
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


class FedstatParser:
    def __init__(self, source_file="assets/charts/fedstat.docx"):
        self.source_file = source_file
        self.output_dir = Path("data")
        self.output_dir.mkdir(exist_ok=True)

        # Setup requests session with retry logic
        self.session = requests.Session()
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
            allowed_methods=["GET", "HEAD"],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)

        # Keywords for filtering - matching against text descriptions
        self.keywords = {
            "inflation": [
                "инфляция",
                "цены",
                "consumer price",
                "ипц",
                "cpi",
                "цены потребителей",
            ],
            "price": ["цена", "price", "тариф", "тарифы", "индекс цен"],
            "production": [
                "производство",
                "production",
                "выпуск",
                "объем",
                "промышленность",
            ],
            "salary": [
                "зарплата",
                "заработная плата",
                "salary",
                "wage",
                "доход",
                "income",
            ],
            "grp": ["врп", "валовой региональный", "GRP", "региональный продукт"],
        }

    def parse_docx(self) -> List[Tuple[str, str]]:
        """Parse docx file and extract (Text, URL) pairs."""
        if not Path(self.source_file).exists():
            raise FileNotFoundError(f"Source file not found: {self.source_file}")

        doc = Document(self.source_file)
        pairs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Extract URLs from the paragraph text
            urls = re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', text)

            for url in urls:
                # Extract the text description (text without the URL)
                description = text.replace(url, "").strip()
                pairs.append((description, url))

        # Also check hyperlinks in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    urls = re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', text)
                    for url in urls:
                        description = text.replace(url, "").strip()
                        pairs.append((description, url))

        return pairs

    def extract_indicator_id(self, url: str) -> Optional[str]:
        """Extract indicator ID from fedstat URL."""
        match = re.search(r"/indicator/(\d+)", url)
        return match.group(1) if match else None

    def categorize_by_keywords(self, text: str) -> List[str]:
        """Categorize based on keyword matching in text."""
        text_lower = text.lower()
        matched_categories = []

        for category, keyword_list in self.keywords.items():
            for keyword in keyword_list:
                if keyword.lower() in text_lower:
                    matched_categories.append(category)
                    break

        return matched_categories if matched_categories else ["other"]

    def calculate_priority_score(self, item: Dict) -> int:
        """Calculate priority score based on keyword matches."""
        categories = item.get("categories", [])

        score = 0

        # Base score by category priority
        category_scores = {
            "inflation": 10,
            "price": 8,
            "salary": 6,
            "production": 5,
            "grp": 4,
            "other": 1,
        }

        # Add scores for all matched categories
        for cat in categories:
            score += category_scores.get(cat, 0)

        # Bonus for specific text patterns
        text_lower = item.get("text", "").lower()
        if any(x in text_lower for x in ["индекс", "index", "динамика", "trend"]):
            score += 1

        return score

    def fetch_indicator_metadata(self, indicator_id: str, url: str) -> Dict:
        """
        Fetch metadata for a specific indicator to verify data availability.

        Args:
            indicator_id: The indicator ID (e.g., "30925")
            url: The full indicator URL

        Returns:
            Dictionary with metadata information
        """
        metadata = {
            "indicator_id": indicator_id,
            "url": url,
            "accessible": False,
            "status_code": None,
            "has_data": False,
            "download_available": False,
            "title": "",
            "comment": "",
            "error": None,
        }

        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }

            # Check page accessibility
            response = self.session.get(
                url, headers=headers, timeout=10, allow_redirects=True
            )
            metadata["status_code"] = response.status_code
            metadata["accessible"] = response.status_code == 200

            if response.status_code == 200:
                html_content = response.text

                # Check for data availability indicators
                # 1. Look for download button/links
                if (
                    "format=excel" in html_content
                    or "format=xlsx" in html_content
                    or "format=csv" in html_content
                ):
                    metadata["download_available"] = True

                # 2. Look for data table indicators (use substring matching)
                data_indicators = [
                    "tabsdata",
                    "table_tabs_content",
                    "fgrid",
                    "onclick=",
                    "tabsdata_item",
                ]
                if any(indicator in html_content for indicator in data_indicators):
                    metadata["has_data"] = True

                # 3. Look for error messages (specific to actual missing data, not UI text)
                error_patterns = ["Данные недоступны", "Data not available"]
                if any(
                    error.lower() in html_content.lower() for error in error_patterns
                ):
                    metadata["has_data"] = False

                # 4. Try to extract title/comment
                if "<title>" in html_content:
                    title_match = re.search(r"<title>([^<]+)</title>", html_content)
                    if title_match:
                        metadata["title"] = title_match.group(1).strip()

                # 5. Look for comment/description
                comment_patterns = [
                    r'<div id="commentIndicator[^"]*">([^<]+)',
                    r'<div class="[^"]*pass_name[^"]*">([^<]+)',
                    r"Комментарии\s*</div>\s*<div[^>]*>([^<]+)",
                ]
                for pattern in comment_patterns:
                    match = re.search(pattern, html_content)
                    if match:
                        metadata["comment"] = match.group(1).strip()
                        break

            else:
                metadata["error"] = f"HTTP {response.status_code}"

        except requests.exceptions.Timeout:
            metadata["error"] = "Timeout"
        except requests.exceptions.RequestException as e:
            metadata["error"] = f"Request error: {str(e)}"
        except Exception as e:
            metadata["error"] = f"Unexpected error: {str(e)}"

        return metadata

    def fetch_metadata_batch(
        self, items: List[Dict], batch_size: int = 5, delay: float = 1.0
    ) -> List[Dict]:
        """
        Fetch metadata for a batch of items with rate limiting.

        Args:
            items: List of items with 'indicator_id' and 'url'
            batch_size: Number of items to process before a pause
            delay: Delay between batches in seconds

        Returns:
            List of items with added metadata
        """
        results = []

        print(f"      Fetching metadata for {len(items)} indicators...")
        print(f"      Batch size: {batch_size}, Delay: {delay}s")

        for i, item in enumerate(items, 1):
            indicator_id = item.get("indicator_id", "")
            url = item.get("url", "")

            if indicator_id:
                metadata = self.fetch_indicator_metadata(indicator_id, url)
                item["metadata"] = metadata

                # Determine overall data availability
                data_available = metadata.get("accessible", False) and metadata.get(
                    "has_data", False
                )
                item["data_available"] = data_available

                status_icon = "✓" if data_available else "✗"
                print(
                    f"      [{i}/{len(items)}] {indicator_id}: {status_icon} "
                    + f"(status: {metadata.get('status_code')}, has_data: {metadata.get('has_data')})"
                )

            results.append(item)

            # Rate limiting: pause between batches
            if i % batch_size == 0 and i < len(items):
                time.sleep(delay)

        return results

    def create_catalog_csv(
        self, pairs: List[Tuple[str, str]], output_path="data/fedstat_catalog.csv"
    ):
        """Create catalog CSV with name-link pairs."""
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["index", "text", "url", "indicator_id", "categories"])

            for i, (text, url) in enumerate(pairs, 1):
                indicator_id = self.extract_indicator_id(url)
                categories = self.categorize_by_keywords(text)
                writer.writerow([i, text, url, indicator_id, "|".join(categories)])

        print(f"Catalog created: {output_path} with {len(pairs)} entries")
        return output_path

    def create_prioritized_feed(
        self, items: List[Dict], output_path="data/prioritized_fedstat_feed.csv"
    ):
        """Create prioritized feed CSV with top candidates including metadata."""
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(
                [
                    "rank",
                    "text",
                    "url",
                    "indicator_id",
                    "categories",
                    "priority_score",
                    "data_available",
                    "accessible",
                    "has_data",
                    "download_available",
                    "status_code",
                    "title",
                    "comment",
                    "error",
                ]
            )

            for i, item in enumerate(items, 1):
                metadata = item.get("metadata", {})
                notes_parts = []

                if metadata.get("accessible"):
                    notes_parts.append(
                        f"Accessible (HTTP {metadata.get('status_code')})"
                    )
                else:
                    notes_parts.append(
                        f"Not accessible (HTTP {metadata.get('status_code', 'N/A')})"
                    )

                if metadata.get("has_data"):
                    notes_parts.append("Has data")
                else:
                    notes_parts.append("No data detected")

                if metadata.get("download_available"):
                    notes_parts.append("Download available")

                if metadata.get("error"):
                    notes_parts.append(f"Error: {metadata.get('error')}")

                notes = "; ".join(notes_parts)

                writer.writerow(
                    [
                        i,
                        item.get("text", ""),
                        item.get("url", ""),
                        item.get("indicator_id", ""),
                        "|".join(item.get("categories", [])),
                        item.get("priority_score", 0),
                        item.get("data_available", False),
                        metadata.get("accessible", False),
                        metadata.get("has_data", False),
                        metadata.get("download_available", False),
                        metadata.get("status_code", ""),
                        metadata.get("title", "")[:100]
                        if metadata.get("title")
                        else "",
                        metadata.get("comment", "")[:200]
                        if metadata.get("comment")
                        else "",
                        metadata.get("error", ""),
                    ]
                )

        print(f"Prioritized feed created: {output_path} with {len(items)} entries")
        return output_path

    def run(self, top_n=50, fetch_metadata=True, batch_size=5):
        """
        Main execution pipeline with optional metadata fetching.

        Args:
            top_n: Number of top candidates to include in output
            fetch_metadata: If True, fetch metadata for top candidates to verify data availability
            batch_size: Number of items to process before a pause during metadata fetching
        """
        print("=" * 60)
        print("Fedstat Smart Link Prioritization")
        print("=" * 60)

        print(f"\n[1/6] Parsing docx file: {self.source_file}...")
        pairs = self.parse_docx()
        print(f"      Extracted {len(pairs)} text-url pairs")

        # Build item list with metadata
        items = []
        for text, url in pairs:
            indicator_id = self.extract_indicator_id(url)
            categories = self.categorize_by_keywords(text)
            items.append(
                {
                    "text": text,
                    "url": url,
                    "indicator_id": indicator_id or "",
                    "categories": categories,
                }
            )

        print(f"\n[2/6] Analyzing keyword matches...")
        category_counts = {}
        for item in items:
            for cat in item["categories"]:
                category_counts[cat] = category_counts.get(cat, 0) + 1

        print(f"      Categories found:")
        for cat, count in sorted(
            category_counts.items(), key=lambda x: x[1], reverse=True
        ):
            print(f"        - {cat}: {count}")

        print(f"\n[3/6] Creating catalog CSV...")
        catalog_path = self.create_catalog_csv(pairs)

        # Verify catalog has >5000 name-link pairs
        if len(pairs) >= 5000:
            print(
                f"      ✓ Catalog has {len(pairs)} name-link pairs (>5000 requirement met)"
            )
        else:
            print(f"      ✗ Catalog has only {len(pairs)} entries (<5000 requirement)")

        print(f"\n[4/6] Calculating priority scores...")
        for item in items:
            item["priority_score"] = self.calculate_priority_score(item)

        # Sort by priority score descending
        items.sort(key=lambda x: (-x["priority_score"], x["text"]))

        # Filter items with score > 0 to get relevant candidates
        relevant_items = [item for item in items if item["priority_score"] > 0]
        print(f"      Found {len(relevant_items)} relevant candidates (score > 0)")

        print(f"\n[5/6] Selecting top {top_n} candidates...")
        top_items = relevant_items[:top_n]

        print(f"      Selected top {len(top_items)} candidates")

        # Fetch metadata for verification (Step 4 of task requirements)
        data_available_count = 0
        api_errors = 0
        if fetch_metadata:
            print(f"\n[5.5/6] Fetching metadata to verify data availability...")
            print(f"      This will make HTTP requests to fedstat.ru")
            top_items = self.fetch_metadata_batch(top_items, batch_size=10, delay=0.2)

            # Count data availability
            for item in top_items:
                if item.get("data_available", False):
                    data_available_count += 1
                if item.get("metadata", {}).get("error"):
                    api_errors += 1

            print(
                f"      Data available: {data_available_count}/{len(top_items)} candidates"
            )
            if api_errors > 0:
                print(f"      API errors encountered: {api_errors}")

        feed_path = self.create_prioritized_feed(top_items)

        print(f"\n[6/6] Output file created")

        print("\n" + "=" * 60)
        print("SUMMARY")
        print("=" * 60)
        print(f"Total name-link pairs in catalog: {len(pairs)}")
        print(f"Relevant candidates found:       {len(relevant_items)}")
        print(f"Top prioritized candidates:      {len(top_items)}")
        if fetch_metadata:
            print(
                f"Data available indicators:       {data_available_count}/{len(top_items)}"
            )
            print(f"API errors:                    {api_errors}")
        print(f"\nOutput files:")
        print(f"  Catalog:   {catalog_path}")
        print(f"  Priority:  {feed_path}")
        print("=" * 60)

        # Display top 10 candidates
        print("\nTop 10 Candidates:")
        print("-" * 60)
        for item in top_items[:10]:
            text_preview = item["text"][:50] if len(item["text"]) > 50 else item["text"]
            data_status = "✓ DATA" if item.get("data_available") else "✗ NO DATA"
            print(f"  [{item['priority_score']:2d}] {text_preview:45s} | {data_status}")

        return {
            "catalog_path": catalog_path,
            "feed_path": feed_path,
            "total_pairs": len(pairs),
            "relevant": len(relevant_items),
            "top": len(top_items),
            "data_available": data_available_count if fetch_metadata else None,
            "api_errors": api_errors if fetch_metadata else None,
        }


if __name__ == "__main__":
    parser = FedstatParser()
    # Run with metadata fetching enabled (task requirement step 4)
    result = parser.run(top_n=50, fetch_metadata=True, batch_size=5)

    print("\n" + "=" * 60)
    print("COMPLETED_TASK")
    print("=" * 60)
