#!/usr/bin/env python3
"""Build April 2026 forecast deviation report artifacts.

Outputs:
- DOCX report ready to send
- XLSX workbook with calculations and reusable template sheets

Important: May 2026 official fact is intentionally not used. Any weekly May rows
in operational files are not monthly facts and are excluded from evidence.
"""
from __future__ import annotations

import csv
import math
import re
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "archive" / "results" / "april_2026_deviation_analysis"
DOCX_PATH = OUT_DIR / "april_2026_forecast_deviation_analysis.docx"
XLSX_PATH = OUT_DIR / "april_2026_forecast_deviation_calculations.xlsx"

FORECAST_INDEX = 100.45
FORECAST_MONTH = "2026-04"
W = {"Prod": 0.3986, "Nonprod": 0.3638, "Serv": 0.2376}
W_WEEKLY = {"food": 0.3986, "nonfood": 0.3638, "services": 0.2376}
COMP_LABELS = {"Prod": "Продовольственные товары", "Nonprod": "Непродовольственные товары", "Serv": "Услуги"}
WEEKLY_COMP_RU = {"food": "Продовольственные", "nonfood": "Непродовольственные", "services": "Услуги"}


def fmt_pct(x: float, digits: int = 2, sign: bool = True) -> str:
    if x is None or (isinstance(x, float) and not math.isfinite(x)):
        return "—"
    prefix = "+" if sign and x > 0 else ""
    return f"{prefix}{x:.{digits}f}%".replace(".", ",")


def fmt_pp(x: float, digits: int = 2, sign: bool = True) -> str:
    if x is None or (isinstance(x, float) and not math.isfinite(x)):
        return "—"
    prefix = "+" if sign and x > 0 else ""
    return f"{prefix}{x:.{digits}f}".replace(".", ",") + " п.п."


def fmt_num(x: float, digits: int = 2) -> str:
    if x is None or (isinstance(x, float) and not math.isfinite(x)):
        return "—"
    return f"{x:.{digits}f}".replace(".", ",")


def norm(s: Any) -> str:
    return re.sub(r"\s+", " ", str(s).strip().lower().replace("ё", "е"))


def parse_num(x: Any) -> float:
    try:
        return float(str(x).strip().replace("\xa0", "").replace(",", "."))
    except Exception:
        return np.nan


def classify_component(x: Any) -> str | None:
    s = norm(x)
    if "непродовольств" in s:
        return "nonfood"
    if "продовольств" in s:
        return "food"
    if "услуг" in s:
        return "services"
    return None


def trim_mean(values: pd.Series, proportion: float = 0.1) -> float:
    arr = np.sort(values.dropna().to_numpy(dtype=float))
    if len(arr) == 0:
        return np.nan
    k = int(np.floor(len(arr) * proportion))
    trimmed = arr[k : len(arr) - k] if len(arr) - 2 * k > 0 else arr
    return float(np.mean(trimmed))


def load_monthly() -> pd.DataFrame:
    df = pd.read_csv(ROOT / "data" / "inflation_data.csv", sep=";", decimal=",", encoding="utf-8-sig")
    df["Date"] = pd.to_datetime(df["Date"], format="%d.%m.%Y")
    for c in ["mom", "Prod", "Nonprod", "Serv", "usd_nom_i", "Ki_i", "Ruonia", "Ruonia_i", "fl_potrb_zad", "fl_dep", "all_real", "Ki"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")
    return df


def build_weekly_product_data() -> tuple[pd.DataFrame, dict[str, str]]:
    hist_path = ROOT / "data" / "kbr_weekly_prices_2008_2026.csv"
    parts: list[pd.DataFrame] = []
    if hist_path.exists():
        hist = pd.read_csv(hist_path).rename(columns={"product_name": "name"})
        hist["date"] = pd.to_datetime(hist["date"])
        hist["price"] = pd.to_numeric(hist["price"], errors="coerce")
        hist["name_norm"] = hist["name"].map(norm)
        hist["source"] = "hist"
        parts.append(hist[["date", "name", "name_norm", "price", "source"]])

    fresh_rows: list[dict[str, Any]] = []
    fresh_path = ROOT / "data" / "Сравнение еженедельных цен_01.csv"
    with fresh_path.open(encoding="utf-8-sig") as f:
        reader = csv.DictReader(f, delimiter=";")
        for row in reader:
            row = {str(k).strip(): v for k, v in row.items()}
            date_raw = str(row.get("Name", "")).strip()
            name = str(row.get("Наименование", "")).strip()
            price = parse_num(row.get("Средние цены, рублей", ""))
            component = classify_component(row.get("Справка_нед.Компоненты", row.get("Компонент", "")))
            wow = parse_num(row.get("Изменение цен, в % к предыдущей неделе", ""))
            if date_raw and name and np.isfinite(price):
                fresh_rows.append(
                    {
                        "date": pd.to_datetime(date_raw, format="%d.%m.%Y", errors="coerce"),
                        "name": name,
                        "name_norm": norm(name),
                        "price": price,
                        "source": "fresh",
                        "component": component,
                        "wow": wow,
                    }
                )
    fresh = pd.DataFrame(fresh_rows).dropna(subset=["date"])
    fresh = fresh.drop_duplicates(["date", "name_norm", "price"])
    comp_map = fresh.dropna(subset=["component"]).drop_duplicates("name_norm").set_index("name_norm")["component"].to_dict()
    parts.append(fresh[["date", "name", "name_norm", "price", "source"]])

    allp = pd.concat(parts, ignore_index=True).dropna(subset=["date", "price"])
    allp["source_rank"] = allp["source"].map({"hist": 0, "fresh": 1})
    allp = allp.sort_values(["date", "name_norm", "source_rank"]).drop_duplicates(["date", "name_norm"], keep="last")
    return allp, comp_map


def compute_analysis() -> dict[str, Any]:
    monthly = load_monthly()
    month_row = monthly[monthly["Date"].dt.strftime("%Y-%m") == FORECAST_MONTH]
    if month_row.empty:
        raise RuntimeError(f"No official monthly fact for {FORECAST_MONTH} in inflation_data.csv")
    apr = month_row.iloc[0]
    may_fact_exists = bool((monthly["Date"].dt.strftime("%Y-%m") == "2026-05").any())

    forecast_mom = FORECAST_INDEX - 100
    fact_index = float(apr["mom"])
    fact_mom = fact_index - 100
    error_pp = forecast_mom - fact_mom

    component_rows = []
    for c in ["Prod", "Nonprod", "Serv"]:
        mom = float(apr[c] - 100)
        contribution = mom * W[c]
        component_rows.append(
            {
                "component": COMP_LABELS[c],
                "code": c,
                "index": float(apr[c]),
                "mom_pct": mom,
                "weight": W[c],
                "contribution_pp": contribution,
            }
        )
    component_df = pd.DataFrame(component_rows)

    aprils = monthly[monthly["Date"].dt.month == 4].copy().sort_values("Date")
    prev_aprils = aprils[aprils["Date"].dt.year < 2026].copy()
    robust_rows = []
    for c, label in [("mom", "Все товары и услуги"), ("Prod", "Продовольственные"), ("Nonprod", "Непродовольственные"), ("Serv", "Услуги")]:
        s = prev_aprils[c] - 100
        x = float(apr[c] - 100)
        q1 = float(s.quantile(0.25))
        q3 = float(s.quantile(0.75))
        iqr = q3 - q1
        lo = q1 - 1.5 * iqr
        hi = q3 + 1.5 * iqr
        outliers = prev_aprils.loc[(s < lo) | (s > hi), "Date"].dt.year.astype(str).tolist()
        s_no = s[(s >= lo) & (s <= hi)]
        median = float(s.median())
        mad = float(np.median(np.abs(s - median)))
        robust_z = (x - median) / (1.4826 * mad) if mad else np.nan
        robust_rows.append(
            {
                "metric": label,
                "april_2026_mom_pct": x,
                "mean_pre2026": float(s.mean()),
                "median_pre2026": median,
                "trim10_pre2026": trim_mean(s),
                "std_pre2026": float(s.std(ddof=1)),
                "q1": q1,
                "q3": q3,
                "iqr": iqr,
                "lower_fence": lo,
                "upper_fence": hi,
                "outlier_years": ", ".join(outliers) if outliers else "нет",
                "mean_no_outliers": float(s_no.mean()),
                "z_score": float((x - s.mean()) / s.std(ddof=1)),
                "robust_z": float(robust_z),
                "percentile_less_than_2026": float((s < x).mean() * 100),
                "min_pre2026": float(s.min()),
                "min_year_pre2026": int(prev_aprils.loc[s.idxmin(), "Date"].year),
            }
        )
    robust_df = pd.DataFrame(robust_rows)

    month_deflation_rows = []
    pre2026 = monthly[monthly["Date"].dt.year < 2026].copy()
    for m, g in pre2026.groupby(pre2026["Date"].dt.month):
        month_deflation_rows.append(
            {
                "month": int(m),
                "n": int(len(g)),
                "deflation_count": int((g["mom"] < 100).sum()),
                "min_mom_pct": float((g["mom"] - 100).min()),
                "median_mom_pct": float((g["mom"] - 100).median()),
                "mean_mom_pct": float((g["mom"] - 100).mean()),
            }
        )
    month_deflation_df = pd.DataFrame(month_deflation_rows)

    allp, comp_map = build_weekly_product_data()
    p0 = allp[allp["date"] == pd.Timestamp("2026-03-30")].set_index("name_norm")
    p1 = allp[allp["date"] == pd.Timestamp("2026-04-27")].set_index("name_norm")
    common = sorted(set(p0.index) & set(p1.index))
    counts = {"food": 0, "nonfood": 0, "services": 0}
    for k in common:
        component = comp_map.get(k)
        if component:
            counts[component] += 1
    product_rows = []
    for k in common:
        component = comp_map.get(k)
        if not component:
            continue
        a = float(p0.loc[k, "price"])
        b = float(p1.loc[k, "price"])
        if a > 0 and b > 0:
            pct = (b / a - 1) * 100
            contribution = pct * W_WEEKLY[component] / counts[component]
            product_rows.append(
                {
                    "product": str(p0.loc[k, "name"]),
                    "component": WEEKLY_COMP_RU[component],
                    "component_code": component,
                    "price_2026_03_30": a,
                    "price_2026_04_27": b,
                    "change_pct": pct,
                    "approx_contribution_pp": contribution,
                }
            )
    product_df = pd.DataFrame(product_rows).sort_values("approx_contribution_pp")

    key_products = [
        "яйца куриные",
        "огурцы свежие",
        "помидоры свежие",
        "картофель",
        "капуста белокочанная",
        "морковь",
        "свекла столовая",
        "лук репчатый",
        "бананы",
    ]
    key_rows = []
    monthly_price_rows = []
    for key in key_products:
        sub = allp[allp["name_norm"].str.contains(norm(key), regex=False)].copy().sort_values("date")
        if sub.empty:
            continue
        sub["month"] = sub["date"].dt.to_period("M")
        mon = sub.groupby("month").tail(1).copy().sort_values("month")
        mon["mom_pct"] = mon["price"].pct_change() * 100
        vals = {str(r["month"]): (float(r["price"]), float(r["mom_pct"]) if pd.notna(r["mom_pct"]) else np.nan) for _, r in mon.iterrows()}

        def ch(a: str, b: str) -> float:
            return (vals[b][0] / vals[a][0] - 1) * 100 if a in vals and b in vals else np.nan

        hist_apr = [ch(f"{y}-03", f"{y}-04") for y in range(2009, 2026) if f"{y}-03" in vals and f"{y}-04" in vals]
        apr2026 = ch("2026-03", "2026-04")
        sep_mar = ch("2025-09", "2026-03")
        sep_apr = ch("2025-09", "2026-04")
        percentile = float(np.mean(np.array(hist_apr) < apr2026) * 100) if hist_apr and np.isfinite(apr2026) else np.nan
        key_rows.append(
            {
                "product_key": key,
                "product_name": str(mon["name"].iloc[-1]),
                "sep2025_to_mar2026_pct": sep_mar,
                "mar2026_to_apr2026_pct": apr2026,
                "sep2025_to_apr2026_pct": sep_apr,
                "historical_april_median_pct": float(np.median(hist_apr)) if hist_apr else np.nan,
                "historical_april_mean_pct": float(np.mean(hist_apr)) if hist_apr else np.nan,
                "april_2026_percentile_vs_history": percentile,
            }
        )
        for _, r in mon[(mon["month"] >= pd.Period("2023-01")) & (mon["month"] <= pd.Period("2026-04"))].iterrows():
            monthly_price_rows.append(
                {
                    "product_key": key,
                    "month": str(r["month"]),
                    "price": float(r["price"]),
                    "mom_pct": float(r["mom_pct"]) if pd.notna(r["mom_pct"]) else np.nan,
                }
            )
    key_df = pd.DataFrame(key_rows)
    key_monthly_df = pd.DataFrame(monthly_price_rows)

    pattern_2022 = monthly[(monthly["Date"] >= "2022-01-01") & (monthly["Date"] <= "2022-12-31")].copy()
    pattern_recent = monthly[(monthly["Date"] >= "2025-09-01") & (monthly["Date"] <= "2026-04-30")].copy()

    return {
        "monthly": monthly,
        "april_fact": apr,
        "may_fact_exists": may_fact_exists,
        "forecast_index": FORECAST_INDEX,
        "forecast_mom": forecast_mom,
        "fact_index": fact_index,
        "fact_mom": fact_mom,
        "error_pp": error_pp,
        "component_df": component_df,
        "aprils": aprils,
        "prev_aprils": prev_aprils,
        "robust_df": robust_df,
        "month_deflation_df": month_deflation_df,
        "product_df": product_df,
        "key_df": key_df,
        "key_monthly_df": key_monthly_df,
        "pattern_2022": pattern_2022,
        "pattern_recent": pattern_recent,
    }


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = fmt_num(val, 3 if abs(val) < 1 else 2)
            else:
                cells[i].text = str(val)


def build_docx(a: dict[str, Any]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Анализ отклонения прогноза ИПЦ КБР за апрель 2026 года")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.add_run("Важно: ").bold = True
    p.add_run("официальный факт за май 2026 года отсутствует. Оперативные недельные наблюдения мая не используются как месячный факт и не включены в доказательную часть анализа.")

    add_heading(doc, "1. Резюме", 1)
    doc.add_paragraph(
        f"Фактический ИПЦ КБР в апреле 2026 года составил {fmt_num(a['fact_index'], 2)}, "
        f"или {fmt_pct(a['fact_mom'])} MoM, при направленном прогнозе {fmt_num(a['forecast_index'], 2)} "
        f"({fmt_pct(a['forecast_mom'])} MoM). Отклонение прогноза от факта составило "
        f"{fmt_pp(a['error_pp'])}, что превышает целевой коридор ±0,5 п.п. на {fmt_num(abs(a['error_pp']) - 0.5, 2)} п.п."
    )
    doc.add_paragraph(
        "Основная причина отклонения — реализация исторически нетипичного дефляционного сценария: "
        "в доступном ряду с 2010 года апрельская дефляция по общему ИПЦ КБР ранее не наблюдалась. "
        "Главный вклад в отклонение внёс продовольственный компонент, прежде всего свежие овощи, бананы и яйца."
    )

    add_heading(doc, "2. План-факт", 1)
    plan_df = pd.DataFrame(
        [
            {"Показатель": "Отправленный прогноз", "Индекс": a["forecast_index"], "MoM": a["forecast_mom"]},
            {"Показатель": "Факт апрель 2026", "Индекс": a["fact_index"], "MoM": a["fact_mom"]},
            {"Показатель": "Отклонение, прогноз − факт", "Индекс": np.nan, "MoM": a["error_pp"]},
        ]
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Показатель", "Индекс", "MoM / п.п."]):
        table.rows[0].cells[i].text = h
    for _, r in plan_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["Показатель"])
        cells[1].text = "—" if pd.isna(r["Индекс"]) else fmt_num(float(r["Индекс"]), 2)
        cells[2].text = fmt_pct(float(r["MoM"])) if "Отклонение" not in r["Показатель"] else fmt_pp(float(r["MoM"]))

    add_heading(doc, "3. Компонентный вклад в MoM", 1)
    comp = a["component_df"].copy()
    comp_show = comp.rename(columns={"component": "Компонент", "index": "Индекс", "mom_pct": "MoM", "weight": "Вес", "contribution_pp": "Вклад"})
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Компонент", "Индекс", "MoM", "Вес", "Вклад в общий MoM"]):
        table.rows[0].cells[i].text = h
    for _, r in comp_show.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["Компонент"])
        cells[1].text = fmt_num(float(r["Индекс"]), 2)
        cells[2].text = fmt_pct(float(r["MoM"]))
        cells[3].text = fmt_num(float(r["Вес"]), 4)
        cells[4].text = fmt_pp(float(r["Вклад"]))
    cells = table.add_row().cells
    cells[0].text = "Итого по компонентам"
    cells[1].text = "—"
    cells[2].text = "—"
    cells[3].text = "—"
    cells[4].text = fmt_pp(float(comp["contribution_pp"].sum()))
    doc.add_paragraph(
        "Продовольственные товары снизились на -0,65% MoM и внесли около -0,26 п.п. в общий индекс. "
        "Непродовольственные товары добавили около -0,06 п.п. вниз. Услуги, напротив, дали небольшой положительный вклад около +0,04 п.п., но не компенсировали товарную дефляцию."
    )

    add_heading(doc, "4. Исторический контекст и устойчивые метрики", 1)
    r_all = a["robust_df"].iloc[0]
    doc.add_paragraph(
        f"В 2010–2025 годах средний апрельский прирост общего ИПЦ составлял {fmt_pct(r_all['mean_pre2026'])}, "
        f"медианный — {fmt_pct(r_all['median_pre2026'])}, усечённое среднее 10% — {fmt_pct(r_all['trim10_pre2026'])}. "
        f"Минимальный апрель до 2026 года составлял {fmt_pct(r_all['min_pre2026'])} в {int(r_all['min_year_pre2026'])} году. "
        f"Апрель 2026 года ({fmt_pct(a['fact_mom'])}) оказался ниже исторического минимума и ниже нижней IQR-границы ({fmt_pct(r_all['lower_fence'])}), "
        f"то есть формально является нижним outlier. Robust z-score относительно медианы/MAD равен {fmt_num(r_all['robust_z'], 2)}."
    )
    robust_small = a["robust_df"][["metric", "april_2026_mom_pct", "mean_pre2026", "median_pre2026", "trim10_pre2026", "outlier_years", "robust_z"]].copy()
    robust_small.columns = ["Показатель", "Апрель 2026", "Среднее", "Медиана", "Усеч. среднее", "Outliers до 2026", "Robust z"]
    table = doc.add_table(rows=1, cols=len(robust_small.columns))
    table.style = "Table Grid"
    for i, h in enumerate(robust_small.columns):
        table.rows[0].cells[i].text = h
    for _, r in robust_small.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(robust_small.columns):
            val = r[col]
            if col in ["Апрель 2026", "Среднее", "Медиана", "Усеч. среднее"]:
                cells[i].text = fmt_pct(float(val))
            elif col == "Robust z":
                cells[i].text = fmt_num(float(val), 2)
            else:
                cells[i].text = str(val)

    doc.add_paragraph(
        "До 2026 года апрель не был дефляционным месяцем в доступном ряду. Дефляционные месяцы чаще появлялись позже — с мая–июня и особенно летом. "
        "В 2026 году часть поздневесенне-летнего дефляционного паттерна проявилась уже в апреле."
    )

    add_heading(doc, "5. Продовольственные драйверы", 1)
    neg = a["product_df"].head(12).copy()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Товар", "Компонент", "Цена 30.03", "Цена 27.04", "Вклад / изменение"]):
        table.rows[0].cells[i].text = h
    for _, r in neg.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["product"])
        cells[1].text = str(r["component"])
        cells[2].text = fmt_num(float(r["price_2026_03_30"]), 2)
        cells[3].text = fmt_num(float(r["price_2026_04_27"]), 2)
        cells[4].text = f"{fmt_pp(float(r['approx_contribution_pp']), 3)}; {fmt_pct(float(r['change_pct']))}"
    doc.add_paragraph(
        "Наиболее сильные отрицательные товарные движения: огурцы (-22,6%), помидоры (-17,9%), бананы (-8,5%), яйца (-7,6%). "
        "Они совпали с коррекцией после заметного зимне-весеннего роста и стали основой продовольственного отрицательного вклада."
    )

    key = a["key_df"]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Товар", "Сен.25–мар.26", "Мар.–апр.26", "Сен.25–апр.26", "Медиана апреля", "Перцентиль апр.26"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for _, r in key.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["product_key"])
        cells[1].text = fmt_pct(float(r["sep2025_to_mar2026_pct"]))
        cells[2].text = fmt_pct(float(r["mar2026_to_apr2026_pct"]))
        cells[3].text = fmt_pct(float(r["sep2025_to_apr2026_pct"]))
        cells[4].text = fmt_pct(float(r["historical_april_median_pct"]))
        cells[5].text = fmt_num(float(r["april_2026_percentile_vs_history"]), 0) + "%"

    add_heading(doc, "6. Почему именно апрель 2026 года дал аномалию", 1)
    for text in [
        "Ранняя плодоовощная коррекция: обычно выраженное дефляционное давление по овощам усиливается в мае–летом, но в 2026 году сильное снижение огурцов и помидоров началось уже в марте–апреле.",
        "Коррекция после перегрева: яйца с сентября 2025 по март 2026 выросли на 46,2%, помидоры — на 109,0%, огурцы — на 51,1%; в апреле начался откат этих цен.",
        "Слабость непродовольственных товаров: компонент снизился на -0,16% MoM и дал около -0,06 п.п. в общий индекс; в недельном контуре снижались телевизоры, смартфоны, отдельные лекарства и товары повседневного спроса.",
        "Валютный фон: индекс usd_nom_i снизился с 94,93 в марте до 92,16 в апреле, что соответствует заметному укреплению рубля / ослаблению валютного давления на торгуемые товары.",
        "Слабость общего реального индикатора: all_real снизился с 100,60 в марте до 96,50 в апреле, что согласуется с меньшей возможностью переноса издержек в цены.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    add_heading(doc, "7. Аналогия с 2022 годом", 1)
    doc.add_paragraph(
        "Ближайшая механическая аналогия — 2022 год: после сильного весеннего ценового шока и последующего укрепления рубля в июне–августе возникла дефляционная коррекция, прежде всего через продовольствие и товары. "
        "В 2026 году масштаб был меньше, но логика похожа: накопленный рост отдельных товаров, укрепление рубля и последующая товарная коррекция. Отличие в том, что в 2026 году дефляционный сдвиг пришёлся уже на апрель, то есть раньше обычного сезонного окна."
    )

    add_heading(doc, "8. Рекомендуемое заключение", 1)
    doc.add_paragraph(
        "Отклонение прогноза объясняется не завышенным базовым апрельским прогнозом, а совпадением нескольких факторов: раннего плодоовощного снижения, коррекции яиц и отдельных продуктов после накопленного роста, слабости непродовольственных товаров и укрепления рубля. "
        "Совокупно это сформировало первый дефляционный апрель в доступной истории региона и привело к факту 99,72 против прогноза 100,45."
    )

    add_heading(doc, "9. Шаблон для будущих аналогичных анализов", 1)
    for text in [
        "Зафиксировать отправленный прогноз, факт, ошибку и прохождение/непрохождение коридора ±0,5 п.п.",
        "Разложить факт по компонентам с весами: продовольствие, непродовольственные товары, услуги.",
        "Сравнить месяц с историей именно этого календарного месяца: среднее, медиана, усечённое среднее, IQR, outliers, минимум/максимум.",
        "Проверить, является ли факт исторической аномалией или укладывается в обычный сезонный диапазон.",
        "Выделить товарные драйверы: вклад, изменение цены, динамика последних 6–12 месяцев, сравнение с исторической сезонностью.",
        "Описать макро- и сопутствующий фон: курс/валютный индекс, ставки, реальные индикаторы, возможное охлаждение спроса.",
        "Сформулировать 3–5 причин отклонения и отделить официальные факты от оперативных недельных сигналов.",
    ]:
        doc.add_paragraph(text, style="List Number")

    doc.save(DOCX_PATH)


def write_df(ws, df: pd.DataFrame, start_row: int = 1, start_col: int = 1, percent_cols: set[str] | None = None) -> None:
    percent_cols = percent_cols or set()
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for j, col in enumerate(df.columns, start_col):
        cell = ws.cell(start_row, j, col)
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    for i, (_, row) in enumerate(df.iterrows(), start_row + 1):
        for j, col in enumerate(df.columns, start_col):
            val = row[col]
            cell = ws.cell(i, j, None if pd.isna(val) else val)
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if isinstance(val, (float, int, np.floating, np.integer)):
                cell.number_format = "0.000"
    for j, col in enumerate(df.columns, start_col):
        width = min(max(len(str(col)) + 2, 12), 45)
        ws.column_dimensions[get_column_letter(j)].width = width


def build_xlsx(a: dict[str, Any]) -> None:
    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("Summary")
    summary = pd.DataFrame(
        [
            ["Прогнозный месяц", FORECAST_MONTH],
            ["Отправленный прогноз, индекс", a["forecast_index"]],
            ["Отправленный прогноз, MoM %", a["forecast_mom"]],
            ["Факт, индекс", a["fact_index"]],
            ["Факт, MoM %", a["fact_mom"]],
            ["Ошибка, п.п. = прогноз - факт", a["error_pp"]],
            ["Майский факт используется?", "нет; официального факта за май нет"],
        ],
        columns=["Показатель", "Значение"],
    )
    write_df(ws, summary)

    ws = wb.create_sheet("Component contributions")
    write_df(ws, a["component_df"])

    ws = wb.create_sheet("April robust history")
    write_df(ws, a["robust_df"])

    ws = wb.create_sheet("April history")
    aprils = a["aprils"][["Date", "mom", "Prod", "Nonprod", "Serv", "usd_nom_i", "Ki_i", "all_real"]].copy()
    aprils["Date"] = aprils["Date"].dt.strftime("%Y-%m")
    for c in ["mom", "Prod", "Nonprod", "Serv"]:
        aprils[c + "_mom_pct"] = aprils[c] - 100
    write_df(ws, aprils)

    ws = wb.create_sheet("Monthly deflation seasonality")
    write_df(ws, a["month_deflation_df"])

    ws = wb.create_sheet("Product drivers April")
    drivers = a["product_df"][["product", "component", "price_2026_03_30", "price_2026_04_27", "change_pct", "approx_contribution_pp"]].copy()
    write_df(ws, drivers)

    ws = wb.create_sheet("Key product summary")
    write_df(ws, a["key_df"])

    ws = wb.create_sheet("Key product 2023-2026")
    write_df(ws, a["key_monthly_df"])

    ws = wb.create_sheet("2022 pattern")
    p22 = a["pattern_2022"][["Date", "mom", "Prod", "Nonprod", "Serv", "usd_nom_i", "Ki_i", "all_real"]].copy()
    p22["Date"] = p22["Date"].dt.strftime("%Y-%m")
    for c in ["mom", "Prod", "Nonprod", "Serv"]:
        p22[c + "_mom_pct"] = p22[c] - 100
    write_df(ws, p22)

    ws = wb.create_sheet("Recent pattern")
    recent = a["pattern_recent"][["Date", "mom", "Prod", "Nonprod", "Serv", "usd_nom_i", "Ki_i", "all_real"]].copy()
    recent["Date"] = recent["Date"].dt.strftime("%Y-%m")
    for c in ["mom", "Prod", "Nonprod", "Serv"]:
        recent[c + "_mom_pct"] = recent[c] - 100
    write_df(ws, recent)

    ws = wb.create_sheet("Future template")
    template = pd.DataFrame(
        [
            ["1", "План-факт", "Внести прогноз, факт, ошибку, коридор"],
            ["2", "Компоненты", "Рассчитать вклад = MoM компонента × вес"],
            ["3", "История месяца", "Среднее, медиана, trim10, IQR, outliers, минимум/максимум"],
            ["4", "Драйверы", "Товарные вклады, динамика последних 6–12 месяцев, сезонность"],
            ["5", "Макрофон", "Валюта, ставки, реальные индикаторы, спрос"],
            ["6", "Итог", "3–5 причин отклонения и урок для следующего прогноза"],
        ],
        columns=["Шаг", "Блок", "Что делать"],
    )
    write_df(ws, template)

    wb.save(XLSX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    analysis = compute_analysis()
    build_docx(analysis)
    build_xlsx(analysis)
    print(DOCX_PATH)
    print(XLSX_PATH)


if __name__ == "__main__":
    main()
