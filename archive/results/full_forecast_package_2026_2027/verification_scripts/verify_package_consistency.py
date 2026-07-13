#!/usr/bin/env python3
"""Verify package contents and basic DOCX/XLSX send-readiness.

Read-only: prints JSON to stdout and does not modify project files.
"""
from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[4]
OUT = ROOT / "archive" / "results" / "full_forecast_package_2026_2027"
APRIL = ROOT / "archive" / "results" / "april_2026_deviation_analysis"
ZIP_PATH = OUT / "sirena_kbr_forecast_package_2026_2027.zip"


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def main() -> None:
    expected_files = {
        "forecast_2026_2027_mom_yoy.csv",
        "forecast_2026_2027_explanation.docx",
        "april_2026_forecast_deviation_analysis.docx",
        "april_2026_forecast_deviation_calculations.xlsx",
        "verification_report.md",
    }
    mismatches: list[str] = []

    with zipfile.ZipFile(ZIP_PATH) as archive:
        infos = archive.infolist()
        actual_files = {info.filename for info in infos}
        zip_entries = [{"filename": info.filename, "file_size": info.file_size} for info in infos]

    if actual_files != expected_files:
        mismatches.append(f"ZIP contents mismatch: actual={sorted(actual_files)}, expected={sorted(expected_files)}")

    forecast_docx = OUT / "forecast_2026_2027_explanation.docx"
    april_docx = APRIL / "april_2026_forecast_deviation_analysis.docx"
    april_xlsx = APRIL / "april_2026_forecast_deviation_calculations.xlsx"
    forecast_text = docx_text(forecast_docx)
    april_text = docx_text(april_docx)

    required_forecast_phrases = [
        "Пояснения терминов для руководства",
        "Nowcast",
        "SA / сезонно скорректированный ряд",
        "Тарифный перенос",
        "дезинфляционным фактором",
        "декабрь 2027 = 104,02",
    ]
    for phrase in required_forecast_phrases:
        if phrase not in forecast_text:
            mismatches.append(f"Forecast DOCX missing phrase: {phrase}")

    required_april_phrases = [
        "официальный факт за май 2026 года отсутствует",
        "99,72",
        "+0,73 п.п.",
        "первый дефляционный апрель",
    ]
    for phrase in required_april_phrases:
        if phrase not in april_text:
            mismatches.append(f"April DOCX missing phrase: {phrase}")

    workbook = load_workbook(april_xlsx, read_only=True, data_only=True)
    required_sheets = {
        "Summary",
        "Component contributions",
        "April robust history",
        "Product drivers April",
        "Future template",
    }
    sheets = set(workbook.sheetnames)
    missing_sheets = required_sheets - sheets
    if missing_sheets:
        mismatches.append(f"April XLSX missing sheets: {sorted(missing_sheets)}")

    result = {
        "zip_path": str(ZIP_PATH.relative_to(ROOT)),
        "zip_entries": zip_entries,
        "forecast_docx_size": forecast_docx.stat().st_size,
        "april_docx_size": april_docx.stat().st_size,
        "april_xlsx_size": april_xlsx.stat().st_size,
        "forecast_docx_has_glossary": "Пояснения терминов для руководства" in forecast_text,
        "april_docx_declares_no_may_fact": "официальный факт за май 2026 года отсутствует" in april_text,
        "april_xlsx_sheets": workbook.sheetnames,
        "mismatches": mismatches,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
